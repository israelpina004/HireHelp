"""DOCX Exporter Service — generates ATS-friendly resume DOCX files.

Converts structured resume JSON from the rewriter into a clean,
professional Word document optimized for ATS readability.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Adds a styled heading with a bottom border."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
        run.font.size = Pt(13 if level == 2 else 18)


def _add_section_divider(doc: Document) -> None:
    """Adds a thin horizontal rule between sections."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    pf = paragraph.paragraph_format
    from docx.oxml.ns import qn
    pBdr = paragraph._p.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pBdr)


def generate_docx(structured_resume: dict) -> io.BytesIO:
    """Renders structured resume data into an ATS-friendly DOCX.

    Args:
        structured_resume: Dict with keys: contact_info, summary,
            experience, education, skills, projects.

    Returns:
        BytesIO buffer containing the DOCX file bytes.
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)

    # --- Contact Header ---
    contact = structured_resume.get("contact_info", {})
    name = contact.get("name", "Candidate")
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(name)
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    contact_parts = []
    for key in ["email", "phone", "linkedin", "location"]:
        val = contact.get(key)
        if val:
            contact_parts.append(val)
    if contact_parts:
        contact_line = doc.add_paragraph(" • ".join(contact_parts))
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_line.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Professional Summary ---
    summary = structured_resume.get("summary")
    if summary:
        _add_section_divider(doc)
        _add_heading(doc, "Professional Summary", level=2)
        p = doc.add_paragraph(summary)
        for run in p.runs:
            run.font.size = Pt(11)

    # --- Experience ---
    experience = structured_resume.get("experience", [])
    if experience:
        _add_section_divider(doc)
        _add_heading(doc, "Experience", level=2)
        for exp in experience:
            title_line = doc.add_paragraph()
            title_run = title_line.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            dates = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
            date_run = title_line.add_run(f"  |  {dates}")
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                for run in bp.runs:
                    run.font.size = Pt(10.5)

    # --- Education ---
    education = structured_resume.get("education", [])
    if education:
        _add_section_divider(doc)
        _add_heading(doc, "Education", level=2)
        for edu in education:
            edu_line = doc.add_paragraph()
            deg_run = edu_line.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
            deg_run.font.bold = True
            deg_run.font.size = Pt(11)
            years = f"{edu.get('start_year', '')} – {edu.get('end_year', '')}"
            yr_run = edu_line.add_run(f"  |  {years}")
            yr_run.font.size = Pt(10)
            yr_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            details = edu.get("details")
            if details:
                det = doc.add_paragraph(details)
                for run in det.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Skills ---
    skills = structured_resume.get("skills", [])
    if skills:
        _add_section_divider(doc)
        _add_heading(doc, "Skills", level=2)
        for skill_group in skills:
            p = doc.add_paragraph(skill_group)
            for run in p.runs:
                run.font.size = Pt(10.5)

    # --- Projects ---
    projects = structured_resume.get("projects", [])
    if projects:
        _add_section_divider(doc)
        _add_heading(doc, "Projects", level=2)
        for proj in projects:
            proj_line = doc.add_paragraph()
            name_run = proj_line.add_run(proj.get("name", ""))
            name_run.font.bold = True
            name_run.font.size = Pt(11)
            url = proj.get("url")
            if url:
                url_run = proj_line.add_run(f"  |  {url}")
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            desc = proj.get("description")
            if desc:
                d = doc.add_paragraph(desc)
                for run in d.runs:
                    run.font.size = Pt(10.5)

    # --- Set narrow margins for ATS-friendliness ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
